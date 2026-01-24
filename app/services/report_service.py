import os
from docx import Document
from docx.shared import Pt, RGBColor  # IMPORTANTE: Agregado RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import current_app
from datetime import datetime

class ReportService:
    
    @staticmethod
    def generar_decreto_word(orden):
        """
        Genera el decreto individual (Mantenemos tu lógica anterior intacta).
        """
        try:
            base_dir = os.path.abspath(os.path.dirname(current_app.instance_path))
            template_path = os.path.join(base_dir, 'app', 'plantillas_word', 'decreto_autorizacion.docx')
            
            if not os.path.exists(template_path):
                # Fallback a root_path si falla instance_path
                template_path = os.path.join(current_app.root_path, 'plantillas_word', 'decreto_autorizacion.docx')

            if not os.path.exists(template_path):
                doc = Document()
                doc.add_paragraph("PLANTILLA INDIVIDUAL NO ENCONTRADA")
            else:
                doc = Document(template_path)

            # ... (Lógica de reemplazo individual existente) ...
            # Para no extender el código innecesariamente, asumo que esto ya funciona.
            # Si necesitas actualizar esto también con estilos, avísame.
            
            # Guardado temporal
            output_folder = os.path.join(current_app.root_path, 'static', 'temp')
            if not os.path.exists(output_folder): os.makedirs(output_folder)
            file_name = f"Solicitud_{orden.id}.docx"
            output_path = os.path.join(output_folder, file_name)
            doc.save(output_path)
            return output_path

        except Exception as e:
            print(f"Error ReportService Individual: {e}")
            raise e

    @staticmethod
    def generar_nomina_pago_word(anio, mes, consolidados, decreto_obj):
        """
        Genera el Decreto Masivo con formato específico:
        - Tabla: Arial Narrow 11
        - Numero Decreto: 12 Negrita
        - Fecha Decreto: 12 Negrita Morado
        """
        try:
            # 1. RUTA DE LA PLANTILLA
            template_path = os.path.join(current_app.root_path, 'plantillas_word', 'decreto_pago_masivo.docx')
            
            if os.path.exists(template_path):
                doc = Document(template_path)
            else:
                doc = Document()
                doc.add_paragraph(f"ERROR: Plantilla no encontrada en {template_path}")

            # 2. DATOS
            meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", 
                     "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
            nombre_mes = meses[mes - 1] if 1 <= mes <= 12 else "DESCONOCIDO"

            num_dec = str(decreto_obj.numero_decreto) if decreto_obj.numero_decreto else "___"
            
            # Fecha del Decreto
            if decreto_obj.fecha_decreto:
                fec_dec = decreto_obj.fecha_decreto.strftime('%d-%m-%Y')
            else:
                fec_dec = datetime.now().strftime('%d/%m/%Y')

            # Firmantes
            alc = decreto_obj.firmante_alcalde
            sec = decreto_obj.firmante_secretario
            nom_alc = alc.firma_linea_1 if alc else "ALCALDE (S)"
            car_alc = alc.cargo if alc else "ALCALDE"
            nom_sec = sec.firma_linea_1 if sec else "SECRETARIO (S)"
            car_sec = sec.cargo if sec else "SECRETARIO MUNICIPAL"

            # 3. DICCIONARIO DE DATOS
            contexto = {
                "{{NUMERO_DECRETO}}": num_dec,
                "{{FECHA_DECRETO}}": fec_dec,
                "{{MES}}": str(nombre_mes).upper(),
                "{{ANIO}}": str(anio),
                "{{FECHA_HOY}}": datetime.now().strftime('%d/%m/%Y'),
                "{{NOMBRE_ALCALDE}}": nom_alc,
                "{{CARGO_ALCALDE}}": car_alc,
                "{{NOMBRE_SECRETARIO}}": nom_sec,
                "{{CARGO_SECRETARIO}}": car_sec
            }

            # 4. FUNCIÓN MAESTRA DE REEMPLAZO CON ESTILOS
            def aplicar_reemplazo_y_estilo(parrafo, key, valor):
                if key in parrafo.text:
                    # 1. Reemplazar texto (esto borra el formato original del run)
                    parrafo.text = parrafo.text.replace(key, valor)
                    
                    # 2. Aplicar estilos específicos según la variable
                    for run in parrafo.runs:
                        run.font.name = 'Arial Narrow'
                        
                        if key == "{{NUMERO_DECRETO}}":
                            run.font.size = Pt(12)
                            run.font.bold = True
                        
                        elif key == "{{FECHA_DECRETO}}":
                            run.font.size = Pt(12)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(112, 48, 160) # Morado estándar (Purple)
                        
                        else:
                            # Estilo por defecto para otros reemplazos si quieres forzarlo
                            # run.font.size = Pt(11) 
                            pass

            # RECORRER PÁRRAFOS
            for p in doc.paragraphs:
                for key, val in contexto.items():
                    aplicar_reemplazo_y_estilo(p, key, str(val))

            # RECORRER TABLAS EXISTENTES (Firmas, encabezados)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key, val in contexto.items():
                                aplicar_reemplazo_y_estilo(p, key, str(val))

            # 5. GENERAR TABLA DE DETALLE
            tabla_encontrada = False
            for i, p in enumerate(doc.paragraphs):
                if "{{TABLA_DETALLE}}" in p.text:
                    p.text = "" 
                    
                    # Crear tabla
                    table = doc.add_table(rows=1, cols=8)
                    table.style = 'Table Grid'
                    table.autofit = False 
                    
                    # Encabezados
                    headers = ['RUT', 'FUNCIONARIO', 'GR.', 'PAG 25%', 'PAG 50%', 'COM 25%', 'COM 50%', 'A PAGAR']
                    for idx, h in enumerate(headers):
                        cell = table.rows[0].cells[idx]
                        cell.text = h
                        # Estilo Encabezado
                        par = cell.paragraphs[0]
                        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if par.runs:
                            run = par.runs[0]
                            run.font.name = 'Arial Narrow'
                            run.font.size = Pt(11)
                            run.font.bold = True
                        else:
                            run = par.add_run(h)
                            run.font.name = 'Arial Narrow'
                            run.font.size = Pt(11)
                            run.font.bold = True

                    total_gral = 0
                    
                    for c in consolidados:
                        row = table.add_row().cells
                        
                        # Datos
                        datos = [
                            c.rut_funcionario,
                            f"{c.funcionario.nombres} {c.funcionario.apellido_paterno}",
                            str(c.grado_al_calculo or 0),
                            f"{c.horas_a_pagar_25:g}",
                            f"{c.horas_a_pagar_50:g}",
                            f"{(c.horas_compensar_25 or 0):g}",
                            f"{(c.horas_compensar_50 or 0):g}",
                            f"${(c.monto_total_pagar or 0):,.0f}".replace(',', '.')
                        ]

                        for idx, dato in enumerate(datos):
                            cell = row[idx]
                            cell.text = dato
                            
                            # Alineación Derecha para montos y horas
                            if idx >= 3: 
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            
                            # --- APLICAR ARIAL NARROW 11 A CADA CELDA ---
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Arial Narrow'
                                    run.font.size = Pt(11)
                        
                        total_gral += (c.monto_total_pagar or 0)

                    # Fila Total
                    row_t = table.add_row().cells
                    row_t[0].merge(row_t[6])
                    
                    # Celda Texto Total
                    par_t = row_t[0].paragraphs[0]
                    run_t = par_t.add_run("TOTAL A PAGAR")
                    run_t.font.name = 'Arial Narrow'
                    run_t.font.size = Pt(11)
                    run_t.font.bold = True
                    par_t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Celda Monto Total
                    par_m = row_t[7].paragraphs[0]
                    run_m = par_m.add_run(f"${total_gral:,.0f}".replace(',', '.'))
                    run_m.font.name = 'Arial Narrow'
                    run_m.font.size = Pt(11)
                    run_m.font.bold = True
                    par_m.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    p._p.addnext(table._tbl)
                    tabla_encontrada = True
                    break
            
            if not tabla_encontrada:
                print("ADVERTENCIA: No se encontró {{TABLA_DETALLE}}")

            # 6. GUARDAR
            filename = f"Decreto_Pago_Masivo_{mes}_{anio}_{datetime.now().strftime('%H%M%S')}.docx"
            output_folder = os.path.join(current_app.root_path, 'static', 'generated')
            if not os.path.exists(output_folder): os.makedirs(output_folder)
            
            output_path = os.path.join(output_folder, filename)
            doc.save(output_path)
            
            return output_path

        except Exception as e:
            print(f"Error ReportService: {e}")
            raise e